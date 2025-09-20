#!/usr/bin/env python3
"""
Obsidian Checker - Cross-Platform GUI
Modern Tkinter interface for Obsidian vault analysis
Works on Windows, macOS, and Linux
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
import sys
import threading
import subprocess
from pathlib import Path
import json
import re
import time
from typing import Optional, Dict, Any, List

# Import core analysis functions directly
try:
    from obsidian_ai_search import ObsidianAISearch
    AI_AVAILABLE = True
except ImportError:
    AI_AVAILABLE = False

# Import for Word document export
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class ObsidianCheckerGUI:
    def __init__(self, root):
        self.root = root
        self.setup_window()
        self.setup_variables()
        self.create_widgets()
        self.check_ai_availability()
        
    def setup_window(self):
        """Configure the main window"""
        self.root.title("🔗 Obsidian Checker")
        self.root.geometry("800x600")
        self.root.minsize(600, 500)
        
        # Configure style for better cross-platform appearance
        self.style = ttk.Style()
        
        # Try to use a modern theme if available
        available_themes = self.style.theme_names()
        if 'clam' in available_themes:
            self.style.theme_use('clam')
        elif 'alt' in available_themes:
            self.style.theme_use('alt')
            
        # Set up window close protocol
        self.root.protocol("WM_DELETE_WINDOW", self.exit_application)
        
        # Set up keyboard shortcuts
        self.root.bind('<Command-q>', lambda e: self.exit_application())  # macOS
        self.root.bind('<Control-q>', lambda e: self.exit_application())  # Windows/Linux
        self.root.bind('<Command-s>', lambda e: self.export_results_dialog())  # macOS export
        self.root.bind('<Control-s>', lambda e: self.export_results_dialog())  # Windows/Linux export
        self.root.bind('<Escape>', lambda e: self.clear_search())  # Clear search with Escape
        
        # Initialize AI search if available
        self.ai_search = None
        if AI_AVAILABLE:
            try:
                self.ai_search = ObsidianAISearch("")
            except Exception as e:
                print(f"Warning: Failed to initialize AI search: {e}")
            
    def setup_variables(self):
        """Initialize GUI variables"""
        self.vault_path = tk.StringVar()
        self.ai_available = tk.BooleanVar()
        self.check_backlinks = tk.BooleanVar(value=True)
        self.use_ai_search = tk.BooleanVar(value=False)
        self.export_results = tk.BooleanVar(value=False)
        self.search_term = tk.StringVar()
        self.running = False
        
    def create_widgets(self):
        """Create and layout all GUI widgets"""
        # Main container with padding
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure grid weights for responsive layout
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        
        # Title
        title_label = ttk.Label(main_frame, text="🔗 Obsidian Checker", 
                               font=('Helvetica', 16, 'bold'))
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 20))
        
        # Vault selection section
        vault_frame = ttk.LabelFrame(main_frame, text="📁 Obsidian Vault", padding="10")
        vault_frame.grid(row=1, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))
        vault_frame.columnconfigure(1, weight=1)
        
        ttk.Label(vault_frame, text="Vault Path:").grid(row=0, column=0, sticky=tk.W, padx=(0, 10))
        
        self.vault_entry = ttk.Entry(vault_frame, textvariable=self.vault_path, width=50)
        self.vault_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(0, 10))
        
        ttk.Button(vault_frame, text="Browse...", 
                  command=self.browse_vault).grid(row=0, column=2)
        
        ttk.Button(vault_frame, text="Auto-find", 
                  command=self.auto_find_vault).grid(row=0, column=3, padx=(5, 0))
        
        # Add Open Obsidian button
        ttk.Button(vault_frame, text="📱 Open Obsidian", 
                  command=self.open_obsidian).grid(row=0, column=4, padx=(5, 0))
        
        # Quick Search section
        search_frame = ttk.LabelFrame(main_frame, text="🔍 Quick Search", padding="10")
        search_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))
        search_frame.columnconfigure(1, weight=1)
        
        ttk.Label(search_frame, text="Search:").grid(row=0, column=0, sticky=tk.W, padx=(0, 10))
        
        self.search_entry = ttk.Entry(search_frame, textvariable=self.search_term, width=40)
        self.search_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(0, 10))
        self.search_entry.bind('<Return>', lambda e: self.quick_search())
        
        ttk.Button(search_frame, text="🔍 Search", 
                  command=self.quick_search).grid(row=0, column=2, padx=(0, 5))
        
        ttk.Button(search_frame, text="Clear", 
                  command=self.clear_search).grid(row=0, column=3)
        
        # Features section
        features_frame = ttk.LabelFrame(main_frame, text="🔧 Analysis Options", padding="10")
        features_frame.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))
        
        ttk.Checkbutton(features_frame, text="Check backlinks and broken links", 
                       variable=self.check_backlinks).grid(row=0, column=0, sticky=tk.W, pady=2)
        
        self.ai_checkbox = ttk.Checkbutton(features_frame, text="AI semantic search and analysis", 
                                          variable=self.use_ai_search)
        self.ai_checkbox.grid(row=1, column=0, sticky=tk.W, pady=2)
        
        ttk.Checkbutton(features_frame, text="Export results to markdown file", 
                       variable=self.export_results).grid(row=2, column=0, sticky=tk.W, pady=2)
        
        # AI status indicator
        self.ai_status_frame = ttk.Frame(features_frame)
        self.ai_status_frame.grid(row=1, column=1, sticky=tk.E, padx=(10, 0))
        
        self.ai_status_label = ttk.Label(self.ai_status_frame, text="", foreground="gray")
        self.ai_status_label.grid(row=0, column=0)
        
        # Control buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=4, column=0, columnspan=3, pady=(0, 10))
        
        self.run_button = ttk.Button(button_frame, text="🚀 Run Analysis", 
                                    command=self.run_analysis, style="Accent.TButton")
        self.run_button.grid(row=0, column=0, padx=(0, 10))
        
        self.stop_button = ttk.Button(button_frame, text="⏹ Stop", 
                                     command=self.stop_analysis, state=tk.DISABLED)
        self.stop_button.grid(row=0, column=1, padx=(0, 10))
        
        ttk.Button(button_frame, text="⚙️ Settings", 
                  command=self.show_settings).grid(row=0, column=2, padx=(0, 10))
        
        ttk.Button(button_frame, text="❓ Help", 
                  command=self.show_help).grid(row=0, column=3, padx=(0, 10))
        
        # Create export menu button
        self.export_var = tk.StringVar(value="📄 Export")
        self.export_button = ttk.Menubutton(button_frame, textvariable=self.export_var,
                                          state=tk.DISABLED)
        self.export_button.grid(row=0, column=4, padx=(0, 10))
        
        # Create export menu
        export_menu = tk.Menu(self.export_button, tearoff=0)
        export_menu.add_command(label="📄 Export as Markdown", command=lambda: self.export_results_dialog('markdown'))
        export_menu.add_command(label="📘 Export as Word Document", command=lambda: self.export_results_dialog('word'))
        export_menu.add_command(label="📊 Export as HTML (Google Docs Ready)", command=lambda: self.export_results_dialog('html'))
        self.export_button['menu'] = export_menu
        
        ttk.Button(button_frame, text="💪 Exit", 
                  command=self.exit_application).grid(row=0, column=5)
        
        # Results section
        results_frame = ttk.LabelFrame(main_frame, text="📊 Results", padding="10")
        results_frame.grid(row=5, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        results_frame.columnconfigure(0, weight=1)
        results_frame.rowconfigure(0, weight=1)
        main_frame.rowconfigure(5, weight=1)
        
        self.results_text = scrolledtext.ScrolledText(results_frame, height=15, width=80,
                                                     wrap=tk.WORD, state=tk.DISABLED)
        self.results_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Progress bar
        self.progress = ttk.Progressbar(main_frame, mode='indeterminate')
        self.progress.grid(row=6, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 5))
        
        # Status bar
        self.status_var = tk.StringVar(value="Ready")
        status_bar = ttk.Label(main_frame, textvariable=self.status_var, relief=tk.SUNKEN)
        status_bar.grid(row=7, column=0, columnspan=3, sticky=(tk.W, tk.E))
        
    def check_ai_availability(self):
        """Check if AI features are available"""
        try:
            if AI_AVAILABLE and self.ai_search:
                # Check if AI environment exists or if AI search is functional
                ai_env_path = Path("obsidian_ai_env")
                if ai_env_path.exists() or self.ai_search.is_available():
                    self.ai_available.set(True)
                    self.ai_status_label.config(text="✅ AI Ready", foreground="green")
                    self.use_ai_search.set(True)  # Enable by default if available
                else:
                    self.ai_available.set(False)
                    self.ai_status_label.config(text="⚠️ AI Not Set Up", foreground="orange")
                    self.ai_checkbox.config(state=tk.DISABLED)
            else:
                self.ai_available.set(False)
                if AI_AVAILABLE:
                    self.ai_status_label.config(text="⚠️ AI Not Set Up", foreground="orange")
                else:
                    self.ai_status_label.config(text="❌ AI Not Available", foreground="red")
                self.ai_checkbox.config(state=tk.DISABLED)
                
        except Exception as e:
            self.ai_available.set(False)
            self.ai_status_label.config(text="❌ AI Error", foreground="red")
            self.ai_checkbox.config(state=tk.DISABLED)
            
    def browse_vault(self):
        """Open file dialog to select vault directory"""
        directory = filedialog.askdirectory(
            title="Select Obsidian Vault Folder",
            initialdir=os.path.expanduser("~")
        )
        if directory:
            self.vault_path.set(directory)
            if self.is_obsidian_vault(directory):
                self.log_message(f"✅ Valid Obsidian vault selected: {directory}")
            else:
                self.log_message(f"⚠️ Selected directory may not be an Obsidian vault (no .obsidian folder found): {directory}")
            self.log_message(f"Selected vault: {directory}")
    def auto_find_vault(self):
        """Automatically find Obsidian vaults"""
        self.log_message("🔍 Searching for Obsidian vaults...")
        
        # Common Obsidian vault locations
        common_paths = [
            os.path.expanduser("~/Documents"),
            os.path.expanduser("~/Desktop"),
            os.path.expanduser("~/Obsidian"),
            os.path.expanduser("~/Documents/Obsidian"),
            os.path.expanduser("~/iCloud Drive (Archive)/Obsidian") if sys.platform == "darwin" else None,
        ]
        
        found_vaults = []
        
        for base_path in common_paths:
            if base_path and os.path.exists(base_path):
                for root, dirs, files in os.walk(base_path):
                    if '.obsidian' in dirs:
                        found_vaults.append(root)
                        
        if found_vaults:
            # Show selection dialog if multiple vaults found
            if len(found_vaults) == 1:
                self.vault_path.set(found_vaults[0])
                self.log_message(f"✅ Found vault: {found_vaults[0]}")
            else:
                self.show_vault_selection(found_vaults)
        else:
            messagebox.showinfo("No Vaults Found", 
                              "No Obsidian vaults found in common locations.\n"
                              "Please use 'Browse...' to select your vault manually.")
            
    def show_vault_selection(self, vaults):
        """Show dialog to select from multiple found vaults"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Select Vault")
        dialog.geometry("500x300")
        dialog.transient(self.root)
        dialog.grab_set()
        
        ttk.Label(dialog, text="Multiple Obsidian vaults found:", 
                 font=('Helvetica', 12, 'bold')).pack(pady=10)
        
        # Listbox with scrollbar
        frame = ttk.Frame(dialog)
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0, 10))
        
        scrollbar = ttk.Scrollbar(frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        listbox = tk.Listbox(frame, yscrollcommand=scrollbar.set)
        listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=listbox.yview)
        
        for vault in vaults:
            listbox.insert(tk.END, vault)
            
        def select_vault():
            selection = listbox.curselection()
            if selection:
                selected_vault = vaults[selection[0]]
                self.vault_path.set(selected_vault)
                self.log_message(f"✅ Selected vault: {selected_vault}")
                dialog.destroy()
                
        ttk.Button(dialog, text="Select", command=select_vault).pack(pady=5)
    
    def open_obsidian(self):
        """Open Obsidian application with the selected vault"""
        vault_path = self.vault_path.get()
        
        if not vault_path:
            messagebox.showwarning("No Vault Selected", 
                                 "Please select an Obsidian vault first before opening Obsidian.")
            return
            
        if not os.path.exists(vault_path):
            messagebox.showerror("Vault Not Found", 
                               f"The selected vault path does not exist:\n{vault_path}")
            return
            
        if not self.is_obsidian_vault(vault_path):
            result = messagebox.askyesno("Not an Obsidian Vault", 
                                       f"The selected path doesn't appear to be an Obsidian vault (no .obsidian folder found).\n\nDo you still want to open Obsidian with this path?\n\nPath: {vault_path}")
            if not result:
                return
        
        try:
            self.log_message(f"🚀 Opening Obsidian with vault: {vault_path}")
            
            # Use subprocess to open Obsidian with the vault
            if sys.platform == "darwin":  # macOS
                subprocess.run(['open', '-a', 'Obsidian', vault_path], check=True)
            elif sys.platform == "win32":  # Windows
                # Try common Windows installation paths
                obsidian_paths = [
                    os.path.expandvars(r"%LOCALAPPDATA%\Obsidian\Obsidian.exe"),
                    os.path.expandvars(r"%APPDATA%\Obsidian\Obsidian.exe"),
                    r"C:\Program Files\Obsidian\Obsidian.exe"
                ]
                
                obsidian_exe = None
                for path in obsidian_paths:
                    if os.path.exists(path):
                        obsidian_exe = path
                        break
                
                if obsidian_exe:
                    subprocess.run([obsidian_exe, vault_path], check=True)
                else:
                    # Fallback: try to open with default application
                    os.startfile(vault_path)
            else:  # Linux
                try:
                    subprocess.run(['obsidian', vault_path], check=True)
                except FileNotFoundError:
                    # Fallback: open directory in default file manager
                    subprocess.run(['xdg-open', vault_path], check=True)
            
            self.log_message("✅ Obsidian launched successfully!")
            messagebox.showinfo("Success", "Obsidian has been launched with your vault!")
            
        except subprocess.CalledProcessError as e:
            error_msg = f"Failed to launch Obsidian. Make sure Obsidian is installed.\n\nError: {str(e)}"
            self.log_message(f"❌ {error_msg}")
            messagebox.showerror("Launch Failed", error_msg)
        except Exception as e:
            error_msg = f"Unexpected error launching Obsidian: {str(e)}"
            self.log_message(f"❌ {error_msg}")
            messagebox.showerror("Error", error_msg)
            
    def run_analysis(self):
        """Run the Obsidian analysis in a separate thread"""
        if not self.vault_path.get():
            messagebox.showerror("Error", "Please select an Obsidian vault first.")
            return
            
        if not os.path.exists(self.vault_path.get()):
            messagebox.showerror("Error", "Selected vault path does not exist.")
            return
            
        # Disable run button and enable stop button
        self.run_button.config(state=tk.DISABLED)
        self.stop_button.config(state=tk.NORMAL)
        self.running = True
        
        # Start progress bar
        self.progress.start()
        self.status_var.set("Running analysis...")
        
        # Clear previous results
        self.clear_results()
        
        # Run analysis in separate thread
        thread = threading.Thread(target=self.run_analysis_thread)
        thread.daemon = True
        thread.start()
        
    def run_analysis_thread(self):
        """Run the actual analysis (called in separate thread)"""
        try:
            self.log_message("🚀 Starting Obsidian vault analysis...")
            self.log_message(f"📁 Vault: {self.vault_path.get()}")
            
            vault_path = self.vault_path.get()
            
            # Check if vault path is valid
            if not vault_path:
                self.log_message("❌ No vault path specified")
                return
                
            if not os.path.exists(vault_path):
                self.log_message("❌ Vault path does not exist")
                return
            
            # Run analysis based on selected options
            analysis_success = True
            
            if self.check_backlinks.get():
                if self.ai_available.get() and self.use_ai_search.get() and self.ai_search:
                    self.log_message("🤖 Using AI-enhanced analysis")
                    try:
                        # Update AI search vault path
                        self.ai_search.vault_path = vault_path
                        
                        # Run AI analysis if available
                        if not self.ai_search.load_cache():
                            self.log_message("🤖 Building AI index first...")
                            self.ai_search.build_index()
                        
                        # Run backlink check with AI enhancement
                        success, message = self.check_backlinks_core(vault_path)
                        if not success:
                            analysis_success = False
                        
                        self.log_message("\n🤖 AI analysis features available for search.")
                        
                    except Exception as e:
                        self.log_message(f"⚠️ AI analysis failed, falling back to standard: {str(e)}")
                        success, message = self.check_backlinks_core(vault_path)
                        if not success:
                            analysis_success = False
                else:
                    self.log_message("🔍 Using standard analysis")
                    success, message = self.check_backlinks_core(vault_path)
                    if not success:
                        analysis_success = False
            
            # Handle export if requested
            if self.export_results.get():
                export_path = f"analysis_results_{Path(vault_path).name}.md"
                self.log_message(f"\n📄 Exporting results to: {export_path}")
                try:
                    current_results = self.results_text.get(1.0, tk.END).strip()
                    if current_results:
                        export_content = self.format_export_content(current_results)
                        with open(export_path, 'w', encoding='utf-8') as f:
                            f.write(export_content)
                        self.log_message(f"✅ Results exported to: {export_path}")
                    else:
                        self.log_message("⚠️ No results to export")
                except Exception as e:
                    self.log_message(f"❌ Export failed: {str(e)}")
                    analysis_success = False
            
            if analysis_success:
                self.log_message("-" * 60)
                self.log_message("✅ Analysis completed successfully!")
            else:
                self.log_message("-" * 60)
                self.log_message("❌ Analysis completed with errors")
                
        except Exception as e:
            self.log_message(f"❌ Error during analysis: {str(e)}")
            
        finally:
            # Re-enable controls
            self.root.after(0, self.analysis_finished)
            
    def analysis_finished(self):
        """Called when analysis is complete (runs on main thread)"""
        self.run_button.config(state=tk.NORMAL)
        self.stop_button.config(state=tk.DISABLED)
        self.progress.stop()
        self.running = False
        self.status_var.set("Analysis complete")
        
    def stop_analysis(self):
        """Stop the running analysis"""
        self.running = False
        self.status_var.set("Stopping analysis...")
        self.log_message("🛑 Analysis stopped by user")
        
    def log_message(self, message):
        """Add message to results text area (thread-safe)"""
        def update_text():
            self.results_text.config(state=tk.NORMAL)
            self.results_text.insert(tk.END, message + "\n")
            self.results_text.see(tk.END)
            self.results_text.config(state=tk.DISABLED)
            self.update_export_button_state()
            
        self.root.after(0, update_text)
        
    def clear_results(self):
        """Clear the results text area"""
        self.results_text.config(state=tk.NORMAL)
        self.results_text.delete(1.0, tk.END)
        self.results_text.config(state=tk.DISABLED)
        self.update_export_button_state()
        
    def update_export_button_state(self):
        """Enable/disable export button based on whether there are results"""
        current_results = self.results_text.get(1.0, tk.END).strip()
        if current_results:
            self.export_button.config(state=tk.NORMAL)
        else:
            self.export_button.config(state=tk.DISABLED)
        
    def show_settings(self):
        """Show settings dialog"""
        messagebox.showinfo("Settings", "Settings dialog coming soon!\n\nFor now, you can modify options in the Analysis Options section.")
        
    def show_help(self):
        """Show help dialog"""
        help_text = """
🔗 Obsidian Checker - Help

This tool analyzes your Obsidian vault for:
• Broken backlinks and references
• Missing files and attachments
• AI-powered semantic search (if enabled)
• Quick content search
• Export capabilities

How to use:
1. Select your Obsidian vault folder (the one containing .obsidian folder)
2. For quick search: Enter a search term and press Enter or click Search
3. For full analysis: Choose options and click 'Run Analysis'

🔍 Quick Search:
• Enter any text to search across your vault
• Uses AI semantic search if enabled
• Press Enter to search or click the Search button
• Press Escape to clear search field

AI Features:
If AI is enabled, you get additional features:
• Semantic concept search
• Similar file detection
• Content analysis

📄 Export Results:
• Click the 'Export' button to save current results
• Choose from Markdown (.md) or Text (.txt) formats
• Automatic filename generation with timestamp
• Formatted output with analysis metadata

⌨️ Keyboard Shortcuts:
• Enter: Perform search (when in search field)
• Escape: Clear search field
• Cmd+S (Mac) / Ctrl+S (PC): Export results
• Cmd+Q (Mac) / Ctrl+Q (PC): Exit application

Supported file types:
• Markdown files (.md)
• All linked attachments

For more information, see the README.md file.
        """
        
        help_dialog = tk.Toplevel(self.root)
        help_dialog.title("Help - Obsidian Checker")
        help_dialog.geometry("500x400")
        help_dialog.transient(self.root)
        
        text_widget = scrolledtext.ScrolledText(help_dialog, wrap=tk.WORD, padx=10, pady=10)
        text_widget.pack(fill=tk.BOTH, expand=True)
        text_widget.insert(tk.END, help_text)
        text_widget.config(state=tk.DISABLED)
        
        ttk.Button(help_dialog, text="Close", 
                  command=help_dialog.destroy).pack(pady=10)
    
    def quick_search(self):
        """Perform quick search in the selected vault"""
        if not self.vault_path.get():
            messagebox.showerror("Error", "Please select an Obsidian vault first.")
            return
            
        if not self.search_term.get().strip():
            messagebox.showwarning("Warning", "Please enter a search term.")
            self.search_entry.focus()
            return
            
        # Clear previous results and show search info
        self.clear_results()
        search_query = self.search_term.get().strip()
        self.log_message(f"🔍 Quick Search: '{search_query}'")
        self.log_message(f"📁 Vault: {self.vault_path.get()}")
        self.log_message("-" * 50)
        
        # Disable search during operation
        self.search_entry.config(state=tk.DISABLED)
        self.run_button.config(state=tk.DISABLED)
        self.stop_button.config(state=tk.NORMAL)
        self.running = True
        self.progress.start()
        self.status_var.set("Searching...")
        
        # Run search in separate thread
        thread = threading.Thread(target=self.quick_search_thread, args=(search_query,))
        thread.daemon = True
        thread.start()
        
    def quick_search_thread(self, search_query):
        """Run the quick search in a separate thread"""
        try:
            vault_path = self.vault_path.get()
            
            if not vault_path:
                self.log_message("❌ No vault path specified")
                return
                
            if not os.path.exists(vault_path):
                self.log_message("❌ Vault path does not exist")
                return
            
            if self.ai_available.get() and self.use_ai_search.get() and self.ai_search:
                self.log_message("🤖 Using AI semantic search...")
                self.log_message("-" * 30)
                
                try:
                    # Update AI search vault path
                    self.ai_search.vault_path = vault_path
                    
                    # Load or build AI index
                    if not self.ai_search.load_cache():
                        self.log_message("🤖 Building AI index first...")
                        self.ai_search.build_index()
                    
                    # Perform AI semantic search
                    results = self.ai_search.semantic_search(search_query)
                    
                    self.log_message(f"\n🤖 AI Concept Search Results for: '{search_query}'")
                    self.log_message("=" * 60)
                    
                    if results:
                        for i, result in enumerate(results, 1):
                            similarity_pct = result['similarity'] * 100
                            self.log_message(f"\n{i}. 📄 {result['file']} (similarity: {similarity_pct:.1f}%)")
                            self.log_message(f"   {result['preview']}")
                    else:
                        self.log_message("❌ No conceptually related content found")
                    
                    self.log_message("=" * 60)
                    
                except Exception as e:
                    self.log_message(f"⚠️ AI search failed, falling back to text search: {str(e)}")
                    success, message = self.search_vault_core(vault_path, search_query)
                    
            else:
                # Use regular text search
                self.log_message("🔍 Using text search...")
                self.log_message("-" * 30)
                
                success, message = self.search_vault_core(vault_path, search_query)
                
            self.log_message("-" * 30)
            self.log_message("✅ Search completed!")
                
        except Exception as e:
            self.log_message(f"❌ Error during search: {str(e)}")
            
        finally:
            # Re-enable controls
            self.root.after(0, self.search_finished)
            
    def search_finished(self):
        """Called when search is complete (runs on main thread)"""
        self.search_entry.config(state=tk.NORMAL)
        self.run_button.config(state=tk.NORMAL)
        self.stop_button.config(state=tk.DISABLED)
        self.progress.stop()
        self.running = False
        self.status_var.set("Search complete")
        
    def clear_search(self):
        """Clear the search term and focus the search entry"""
        self.search_term.set("")
        self.search_entry.focus()
        
    def export_results_dialog(self, format_type='markdown'):
        """Show export dialog and save results in specified format"""
        # Check if there are results to export
        current_results = self.results_text.get(1.0, tk.END).strip()
        if not current_results:
            messagebox.showwarning("No Results", "No results to export. Please run an analysis or search first.")
            return
            
        # Get default filename based on vault name and format
        vault_name = "results"
        if self.vault_path.get():
            vault_name = Path(self.vault_path.get()).name
        
        # Set file extension and dialog options based on format
        if format_type == 'word':
            if not DOCX_AVAILABLE:
                messagebox.showerror("Word Export Unavailable", 
                                   "Word document export requires python-docx.\n\nInstall it with: pip install python-docx")
                return
            default_extension = ".docx"
            file_types = [("Word Documents", "*.docx"), ("All files", "*.*")]
            title = "Export Results as Word Document"
        elif format_type == 'html':
            default_extension = ".html"
            file_types = [("HTML files", "*.html"), ("All files", "*.*")]
            title = "Export Results as HTML (Google Docs Ready)"
        else:  # markdown (default)
            default_extension = ".md"
            file_types = [("Markdown files", "*.md"), ("Text files", "*.txt"), ("All files", "*.*")]
            title = "Export Results as Markdown"
            
        default_filename = f"obsidian_analysis_{vault_name}_{self.get_timestamp()}{default_extension}"
        
        # Show save dialog
        file_path = filedialog.asksaveasfilename(
            title=title,
            defaultextension=default_extension,
            filetypes=file_types,
            initialfile=default_filename
        )
        
        if file_path:
            try:
                if format_type == 'word':
                    self.export_to_word(current_results, file_path)
                elif format_type == 'html':
                    self.export_to_html(current_results, file_path)
                else:  # markdown
                    self.export_to_markdown(current_results, file_path)
                    
                # Show success message
                format_name = format_type.title() if format_type != 'html' else 'HTML'
                messagebox.showinfo(
                    "Export Successful", 
                    f"Results exported successfully as {format_name} to:\n{file_path}"
                )
                
                self.log_message(f"\n📄 Results exported as {format_name} to: {file_path}")
                
            except Exception as e:
                messagebox.showerror("Export Error", f"Failed to export results:\n{str(e)}")
                self.log_message(f"\n❌ Export failed: {str(e)}")
                
    def export_to_markdown(self, results_text, file_path):
        """Export results to Markdown format"""
        export_content = self.format_export_content(results_text)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(export_content)
    
    def export_to_word(self, results_text, file_path):
        """Export results to Word document format"""
        import datetime
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Obsidian Checker Analysis Results', 0)
        
        # Add metadata
        doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Vault: {self.vault_path.get() or 'Not specified'}")
        doc.add_paragraph(f"Analysis Type: {'AI-Enhanced' if (self.ai_available.get() and self.use_ai_search.get()) else 'Standard'}")
        
        # Add separator
        doc.add_paragraph("_" * 50)
        
        # Process results text and add to document
        lines = results_text.split('\n')
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_paragraph:
                    current_paragraph = None
                continue
                
            # Check for headers/sections
            if line.startswith('=') and len(set(line)) == 1:
                # Skip separator lines
                continue
            elif any(line.startswith(prefix) for prefix in ['📊', '🗺️', '🔍', '🤖']):
                # This looks like a header
                doc.add_heading(line, level=2)
                current_paragraph = None
            elif line.startswith('📄') or line.startswith('   '):
                # This looks like a sub-item or indented content
                if not current_paragraph:
                    current_paragraph = doc.add_paragraph()
                current_paragraph.add_run(line + '\n')
            else:
                # Regular paragraph
                doc.add_paragraph(line)
                current_paragraph = None
        
        # Add footer
        doc.add_paragraph("_" * 50)
        footer = doc.add_paragraph("Generated by Obsidian Checker GUI")
        footer.italic = True
        
        # Save document
        doc.save(file_path)
    
    def export_to_html(self, results_text, file_path):
        """Export results to HTML format (Google Docs ready)"""
        import datetime
        import html
        
        # HTML template
        html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Obsidian Checker Analysis Results</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34495e;
            margin-top: 30px;
            border-left: 4px solid #3498db;
            padding-left: 10px;
        }}
        .metadata {{
            background-color: #f8f9fa;
            padding: 15px;
            border-radius: 5px;
            margin: 20px 0;
        }}
        .results {{
            background-color: #ffffff;
            padding: 20px;
            border: 1px solid #dee2e6;
            border-radius: 5px;
            font-family: 'Monaco', 'Menlo', 'Consolas', monospace;
            white-space: pre-wrap;
        }}
        .footer {{
            margin-top: 30px;
            padding-top: 20px;
            border-top: 1px solid #dee2e6;
            font-style: italic;
            color: #6c757d;
        }}
    </style>
</head>
<body>
    <h1>🔗 Obsidian Checker Analysis Results</h1>
    
    <div class="metadata">
        <p><strong>Generated:</strong> {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        <p><strong>Vault:</strong> {html.escape(self.vault_path.get() or 'Not specified')}</p>
        <p><strong>Analysis Type:</strong> {'AI-Enhanced' if (self.ai_available.get() and self.use_ai_search.get()) else 'Standard'}</p>
    </div>
    
    <h2>Analysis Results</h2>
    <div class="results">{html.escape(results_text)}</div>
    
    <div class="footer">
        <p>Generated by Obsidian Checker GUI</p>
        <p>To import into Google Docs: Open Google Docs → File → Import → Upload this HTML file</p>
    </div>
</body>
</html>
"""
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
    
    def format_export_content(self, results_text):
        """Format the results for markdown export"""
        import datetime
        
        # Header information
        header = f"""
# Obsidian Checker Analysis Results

**Generated:** {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
**Vault:** {self.vault_path.get() or 'Not specified'}
**Analysis Type:** {'AI-Enhanced' if (self.ai_available.get() and self.use_ai_search.get()) else 'Standard'}

---

"""
        
        # Format the results text
        formatted_results = results_text.replace('\n', '\n')
        
        # Footer
        footer = f"""

---

*Generated by Obsidian Checker GUI*
*For more information, visit: https://github.com/your-repo*
"""
        
        return header + formatted_results + footer
        
    def get_timestamp(self):
        """Get current timestamp for filenames"""
        import datetime
        return datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    
    # Core analysis functions - moved from CLI module
    def is_obsidian_vault(self, path):
        """Check if a directory is an Obsidian vault"""
        obsidian_config = os.path.join(path, ".obsidian")
        return os.path.exists(obsidian_config) and os.path.isdir(obsidian_config)
    
    def check_backlinks_core(self, vault_path):
        """Core backlink checking functionality"""
        if not vault_path or not os.path.exists(vault_path):
            return False, "Please provide a valid Obsidian vault directory"
            
        if not self.is_obsidian_vault(vault_path):
            return False, "Selected directory is not an Obsidian vault"
            
        self.log_message(f"🔍 Scanning vault: {vault_path}")
        self.log_message("-" * 60)
        
        try:
            # Find all markdown files
            md_files = list(Path(vault_path).rglob("*.md"))
            total_files = len(md_files)
            
            self.log_message(f"📁 Found {total_files} markdown files")
            
            # Get all file names (without extension) for reference
            all_notes = {f.stem for f in md_files}
            
            broken_links = []
            broken_count = 0
            total_links = 0
            
            for i, md_file in enumerate(md_files):
                if not self.running:
                    return False, "Analysis stopped by user"
                    
                if i % 10 == 0:  # Progress indicator
                    self.log_message(f"📊 Progress: {i+1}/{total_files} files processed...")
                
                try:
                    with open(md_file, 'r', encoding='utf-8') as f:
                        content = f.read()
                    
                    # Find all wiki-style links [[link]]
                    wiki_links = re.findall(r'\[\[([^\]]+)\]\]', content)
                    
                    # Find all markdown links [text](link)
                    md_links = re.findall(r'\[([^\]]*)\]\(([^)]+)\)', content)
                    
                    for link in wiki_links:
                        total_links += 1
                        # Handle links with aliases [[link|alias]]
                        actual_link = link.split('|')[0].strip()
                        
                        # Check if the target note exists
                        if actual_link not in all_notes:
                            # Check if it's a file with extension
                            target_path = Path(vault_path) / f"{actual_link}.md"
                            if not target_path.exists():
                                broken_links.append({
                                    'file': str(md_file.relative_to(vault_path)),
                                    'link': link,
                                    'type': 'wiki'
                                })
                                broken_count += 1
                    
                    for text, link in md_links:
                        total_links += 1
                        # Only check local markdown links
                        if link.endswith('.md') and not link.startswith(('http', 'https', 'ftp')):
                            target_path = md_file.parent / link
                            if not target_path.exists():
                                broken_links.append({
                                    'file': str(md_file.relative_to(vault_path)),
                                    'link': link,
                                    'type': 'markdown'
                                })
                                broken_count += 1
                                
                except Exception as e:
                    self.log_message(f"❌ Error reading {md_file.name}: {str(e)}")
                    
            # Display results
            self.log_message("\n" + "=" * 60)
            self.log_message("📊 BACKLINK CHECK SUMMARY")
            self.log_message("=" * 60)
            self.log_message(f"Files scanned: {total_files}")
            self.log_message(f"Total links found: {total_links}")
            self.log_message(f"Broken links: {broken_count}")
            
            if broken_count == 0:
                self.log_message("\n🎉 All backlinks are working correctly!")
            else:
                self.log_message(f"\n⚠️  Found {broken_count} broken links:")
                self.log_message("-" * 40)
                
                for broken_link in broken_links:
                    link_type = "[[...]]" if broken_link['type'] == 'wiki' else "[...](…)"
                    self.log_message(f"📄 {broken_link['file']}")
                    self.log_message(f"   🔗 {link_type}: {broken_link['link']}")
                    self.log_message("")
                    
            self.log_message("=" * 60)
            return broken_count == 0, f"Analysis completed. {broken_count} broken links found."
            
        except Exception as e:
            error_msg = f"Error during backlink check: {str(e)}"
            self.log_message(f"❌ {error_msg}")
            return False, error_msg
    
    def search_vault_core(self, vault_path, search_term, case_sensitive=False, whole_word=False, use_regex=False):
        """Core search functionality"""
        if not vault_path or not os.path.exists(vault_path):
            return False, "Please provide a valid Obsidian vault directory"
            
        if not self.is_obsidian_vault(vault_path):
            return False, "Selected directory is not an Obsidian vault"
        
        if not search_term.strip():
            return False, "Please provide a search term"
            
        self.log_message(f"🔍 Searching for '{search_term}' in vault: {vault_path}")
        self.log_message("-" * 60)
        
        try:
            # Find all markdown files
            md_files = list(Path(vault_path).rglob("*.md"))
            total_files = len(md_files)
            
            self.log_message(f"📁 Scanning {total_files} markdown files...")
            
            # Prepare search pattern
            if use_regex:
                try:
                    flags = 0 if case_sensitive else re.IGNORECASE
                    pattern = re.compile(search_term, flags)
                except re.error as e:
                    error_msg = f"Invalid regex pattern: {e}"
                    self.log_message(f"❌ {error_msg}")
                    return False, error_msg
            else:
                # Escape special regex characters for literal search
                escaped_term = re.escape(search_term)
                if whole_word:
                    escaped_term = r'\b' + escaped_term + r'\b'
                flags = 0 if case_sensitive else re.IGNORECASE
                pattern = re.compile(escaped_term, flags)
            
            search_results = []
            total_matches = 0
            files_with_matches = 0
            
            for i, md_file in enumerate(md_files):
                if not self.running:
                    return False, "Search stopped by user"
                    
                if i % 10 == 0:  # Progress indicator
                    self.log_message(f"📊 Progress: {i+1}/{total_files} files processed...")
                
                try:
                    with open(md_file, 'r', encoding='utf-8') as f:
                        lines = f.readlines()
                    
                    file_matches = []
                    for line_num, line in enumerate(lines, 1):
                        matches = list(pattern.finditer(line))
                        if matches:
                            file_matches.append({
                                'line_num': line_num,
                                'line_content': line.rstrip(),
                                'matches': len(matches)
                            })
                    
                    if file_matches:
                        files_with_matches += 1
                        file_total_matches = sum(m['matches'] for m in file_matches)
                        total_matches += file_total_matches
                        
                        search_results.append({
                            'file_path': md_file,
                            'relative_path': str(md_file.relative_to(vault_path)),
                            'matches': file_matches,
                            'total_matches': file_total_matches
                        })
                        
                except Exception as e:
                    self.log_message(f"❌ Error reading {md_file.name}: {str(e)}")
            
            # Display results
            self.log_message("\n" + "=" * 60)
            self.log_message(f"📊 SEARCH RESULTS FOR: '{search_term}'")
            self.log_message("=" * 60)
            self.log_message(f"Files scanned: {total_files}")
            self.log_message(f"Files with matches: {files_with_matches}")
            self.log_message(f"Total matches: {total_matches}")
            
            if total_matches == 0:
                self.log_message(f"\n❌ No matches found for '{search_term}'")
            else:
                self.log_message(f"\n✅ Found {total_matches} matches in {files_with_matches} files:")
                self.log_message("-" * 60)
                
                for result in search_results:
                    self.log_message(f"\n📄 {result['relative_path']} ({result['total_matches']} matches)")
                    
                    # Show up to 5 matches per file in GUI
                    for i, match in enumerate(result['matches'][:5]):
                        line_preview = match['line_content'][:100] + "..." if len(match['line_content']) > 100 else match['line_content']
                        self.log_message(f"   Line {match['line_num']}: {line_preview}")
                    
                    if len(result['matches']) > 5:
                        self.log_message(f"   ... and {len(result['matches']) - 5} more matches")
            
            self.log_message("=" * 60)
            return len(search_results) > 0, f"Search completed. {total_matches} matches found."
            
        except Exception as e:
            error_msg = f"Error during search: {str(e)}"
            self.log_message(f"❌ {error_msg}")
            return False, error_msg
    
    def exit_application(self):
        """Exit the application with confirmation"""
        if self.running:
            result = messagebox.askyesno(
                "Exit Application", 
                "An analysis or search is currently running.\n"
                "Do you want to stop it and exit?"
            )
            if result:
                self.running = False
                self.root.after(100, self.root.quit)  # Small delay to stop operations
        else:
            result = messagebox.askyesno(
                "Exit Application", 
                "Are you sure you want to exit Obsidian Checker?"
            )
            if result:
                self.root.quit()

def main():
    """Main function to run the GUI application"""
    root = tk.Tk()
    app = ObsidianCheckerGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()